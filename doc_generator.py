from docx import Document
import io

def generate_checklist_docx(tender_title: str, checklist_data: str) -> io.BytesIO:
    """Compiles extracted operational criteria into a downloadable Word Document."""
    doc = Document()
    
    doc.add_heading("Bid Preparation Checklist", level=1)
    doc.add_paragraph(f"Generated for Source File: {tender_title}")
    
    doc.add_heading("Action Items, Mandates, & Target Timelines", level=2)
    
    # Simple formatting loop to split raw string output into document paragraphs
    for line in checklist_data.split('\n'):
        if line.strip().startswith(('-', '*', '1.', '2.', '3.')):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io